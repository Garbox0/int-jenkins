import tkinter as tk
from tkinter import messagebox, filedialog
from fpdf import FPDF
from docx import Document
import fitz  # PyMuPDF
import re
from PIL import Image, ImageTk

# Funciones para generar archivos
def generate_pdf(data, file_path, template_text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    filled_text = template_text
    for field, value in data.items():
        filled_text = filled_text.replace(field, value)

    for line in filled_text.split('\n'):
        pdf.cell(200, 10, txt=line, ln=True)

    pdf.output(file_path)
    messagebox.showinfo("Success", "PDF generado exitosamente.")
    root.destroy()

def generate_word(data, file_path, template_text):
    doc = Document()

    filled_text = template_text
    for field, value in data.items():
        filled_text = filled_text.replace(field, value)

    for line in filled_text.split('\n'):
        doc.add_paragraph(line)

    doc.save(file_path)
    messagebox.showinfo("Success", "Documento Word generado exitosamente.")
    root.destroy()

# Funciones para leer archivos y encontrar campos vacíos con contexto
def read_pdf(file_path, context_words=1):
    doc = fitz.open(file_path)
    empty_fields = []
    template_text = ""
    field_coords = {}
    field_contexts = {}

    pattern = r"\[[^\]]+\]|-{2,}|\.{2,}"

    field_count = 0
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text = page.get_text("text")
        template_text += text + '\n'
        lines = text.split('\n')

        for line in lines:
            matches = handle_multiple_empty_fields(line)
            for match in matches:
                field = match
                field_count += 1
                unique_field = f"__field_{field_count}__"
                template_text = template_text.replace(field, unique_field)
                field_index = line.index(field)
                context_start = max(0, field_index - context_words)
                context_end = min(len(line), field_index + len(field))
                context = line[context_start:context_end]
                empty_fields.append((unique_field, context))
                field_contexts[unique_field] = (line, page_num)

                rects = page.search_for(field)
                if rects:
                    field_coords[unique_field] = (rects[0], page_num)

    return empty_fields, template_text, field_coords, field_contexts

def read_word(file_path, context_words=1):
    doc = Document(file_path)
    empty_fields = []
    template_text = ""
    field_contexts = {}

    pattern = r"\[[^\]]+\]|-{2,}|\.{2,}"

    field_count = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text
        template_text += text + '\n'
        lines = text.split('\n')

        for line in lines:
            matches = handle_multiple_empty_fields(line)
            for match in matches:
                field = match
                field_count += 1
                unique_field = f"__field_{field_count}__"
                template_text = template_text.replace(field, unique_field)
                field_index = line.index(field)
                context_start = max(0, field_index - context_words)
                context_end = min(len(line), field_index + len(field))
                context = line[context_start:context_end]
                empty_fields.append((unique_field, context))
                field_contexts[unique_field] = (line, paragraph)

    return empty_fields, template_text, field_coords, field_contexts

# Función para manejar múltiples campos vacíos en una línea
def handle_multiple_empty_fields(line):
    pattern = re.compile(r'\[.*?\]|\-{2,}|\.{2,}')
    return pattern.findall(line)

# Función para crear campos de entrada en la GUI
def create_entry_fields(empty_fields, field_coords):
    entries = []
    for i, (field, context) in enumerate(empty_fields):
        label_text = get_relevant_context(field)
        
        if field in field_coords:
            rect, page_num = field_coords[field]
            page = pdf_document.load_page(page_num)
            pix = page.get_pixmap(clip=rect)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            img_tk = ImageTk.PhotoImage(img)
            img_label = tk.Label(root, image=img_tk)
            img_label.image = img_tk
            img_label.grid(row=i, column=0, padx=10, pady=5)

        entry = tk.Entry(root, width=40)
        entry.grid(row=i, column=1, padx=10, pady=5)
        entries.append((field, entry, label_text))
    
    return entries

# Función para extraer solo la palabra anterior al campo vacío
def get_relevant_context(field):
    if field in field_contexts:
        context_text, _ = field_contexts[field]
        pattern = re.compile(r'\[.*?\]|\-{2,}|\.{2,}')
        matches = list(pattern.finditer(context_text))
        for match in matches:
            if match.group(0) == field:
                start = match.start()
                before_text = context_text[:start].strip().split(' ')[-1]  # Obtener la palabra anterior
                return f"{before_text} **{field}**"
    return "Descripción no disponible."

# Función para enviar el formulario
def submit_form(template_text):
    data = {field: entry.get() for field, entry, _ in entries}
    
    if var.get() == 1:
        file_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")])
        if file_path:
            generate_pdf(data, file_path, template_text)
    else:
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if file_path:
            generate_word(data, file_path, template_text)

# Función para seleccionar y leer un archivo
def select_file():
    global pdf_document, field_coords, empty_fields_dict, field_contexts
    file_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf"), ("Word files", "*.docx")])
    if file_path.endswith(".pdf"):
        pdf_document = fitz.open(file_path)
        empty_fields, template_text, field_coords, field_contexts = read_pdf(file_path)
    else:
        empty_fields, template_text, field_coords, field_contexts = read_word(file_path)
    
    empty_fields_dict = dict(empty_fields)

    if empty_fields:
        global entries
        for widget in root.winfo_children():
            widget.grid_forget()
        entries = create_entry_fields(empty_fields, field_coords)
        tk.Radiobutton(root, text="Exportar como PDF", variable=var, value=1).grid(row=len(entries)+1, column=0, padx=10, pady=10)
        tk.Radiobutton(root, text="Exportar como Word", variable=var, value=2).grid(row=len(entries)+1, column=1, padx=10, pady=10)
        tk.Button(root, text="Enviar", command=lambda: submit_form(template_text)).grid(row=len(entries)+2, columnspan=2, pady=20)
    else:
        messagebox.showinfo("Campos Vacíos", "No se encontraron campos vacíos.")

# Crear la ventana principal
root = tk.Tk()
root.title("Formulario de Datos")

var = tk.IntVar()
entries = []
pdf_document = None
field_coords = {}
empty_fields_dict = {}
field_contexts = {}

tk.Button(root, text="Seleccionar Archivo", command=select_file).grid(row=0, columnspan=2, pady=20)

# Iniciar el bucle principal de la ventana
root.mainloop()
